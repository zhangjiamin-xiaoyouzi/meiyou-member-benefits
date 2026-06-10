#!/usr/bin/env python3
"""美柚中后台 PRD 文档生成器"""

from __future__ import annotations

import argparse
import io
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("请安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)

FONT_NAME = "微软雅黑"
TITLE_SIZE = Pt(22)
H1_SIZE = Pt(16)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)
BODY_SIZE = Pt(11)
TABLE_SIZE = Pt(10)
COLOR_TITLE = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_H1 = RGBColor(0x33, 0x33, 0x33)
COLOR_H2 = RGBColor(0x44, 0x44, 0x44)
COLOR_BODY = RGBColor(0x55, 0x55, 0x55)
TABLE_HEADER_BG = "D9E2F3"
TABLE_BORDER = "999999"


def set_cell_shading(cell, color):
    shading = parse_xml('<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        '<w:tblPr {}/>'.format(nsdecls("w")))
    borders = parse_xml(
        '<w:tblBorders {}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="{}"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="{}"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{}"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="{}"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{}"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{}"/>'
        '</w:tblBorders>'.format(
            nsdecls("w"), TABLE_BORDER, TABLE_BORDER,
            TABLE_BORDER, TABLE_BORDER, TABLE_BORDER, TABLE_BORDER))
    tblPr.append(borders)


def add_p(doc, text, font_size=BODY_SIZE, bold=False, color=COLOR_BODY,
          alignment=None, space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = font_size
    run.bold = bold
    run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    return p


def add_table(doc, headers, rows):
    if not rows:
        rows = [[""] * len(headers)]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = TABLE_SIZE
        run.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        set_cell_shading(cell, TABLE_HEADER_BG)
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text) if cell_text else "")
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            run.font.size = TABLE_SIZE
            run.font.color.rgb = COLOR_BODY
    doc.add_paragraph()
    return table


def add_meta_table(doc, meta):
    """添加元数据信息表"""
    fields = [
        ("状态", meta.get("status", "规划中")),
        ("创建模板", meta.get("template_name", "中后台需求模版")),
        ("分类", meta.get("category", "")),
        ("业务", meta.get("business", "")),
        ("迭代", meta.get("iteration", "")),
        ("处理人", meta.get("handler", "")),
        ("优先级", meta.get("priority", "Middle")),
        ("需求类型", meta.get("requirement_type", "")),
        ("需求难度", meta.get("difficulty", "")),
        ("预估工时", meta.get("estimated_hours", "")),
    ]
    rows = [[k, v] for k, v in fields if v]
    add_table(doc, ["字段", "内容"], rows)


def generate_prd(data, output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ── 标题 ──
    title = data.get("title", "{角色}可以通过{功能}获得{价值}")
    add_p(doc, title, font_size=TITLE_SIZE, bold=True, color=COLOR_TITLE,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

    add_p(doc, "ID: " + data.get("story_id", ""), font_size=Pt(10),
          color=RGBColor(0x99, 0x99, 0x99),
          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ── 元数据 ──
    add_meta_table(doc, data.get("meta", {}))

    # ══════ 一、需求背景 ══════
    add_p(doc, "一、需求背景（20%）", font_size=H1_SIZE, bold=True,
          color=COLOR_H1, space_after=Pt(8))

    add_p(doc, "1.1 业务大背景（5%）", font_size=H2_SIZE, bold=True, color=COLOR_H2)
    add_p(doc, data.get("bg_big", "对齐团队 OKR，需求服务于团队哪个大 O…"))

    add_p(doc, "1.2 业务子背景（5%）", font_size=H2_SIZE, bold=True, color=COLOR_H2)
    add_p(doc, data.get("bg_sub", "用户反馈 / Case 分析 / 商业化诉求…"))

    add_p(doc, "1.3 现状判断及问题（10%）", font_size=H2_SIZE, bold=True, color=COLOR_H2)
    status_rows = data.get("status_analysis", [])
    if status_rows:
        add_table(doc, ["现状", "问题判断", "历史需求或复盘文档", "解决方案"],
                  [[r.get("current", ""), r.get("problem", ""),
                    r.get("history", ""), r.get("solution", "")]
                   for r in status_rows])

    # ══════ 二、项目目标 ══════
    add_p(doc, "二、项目目标（5%）", font_size=H1_SIZE, bold=True,
          color=COLOR_H1, space_after=Pt(8))

    add_p(doc, "2.1 目标描述（5%）", font_size=H2_SIZE, bold=True, color=COLOR_H2)
    add_p(doc, data.get("goal", "本次迭代的价值…"))

    if data.get("iteration_plan"):
        add_p(doc, "2.2 迭代节奏", font_size=H2_SIZE, bold=True, color=COLOR_H2)
        add_p(doc, data["iteration_plan"])

    if data.get("risk"):
        add_p(doc, "2.3 风险预判", font_size=H2_SIZE, bold=True, color=COLOR_H2)
        add_p(doc, data["risk"])

    # ══════ 三、需求方案 ══════
    add_p(doc, "三、需求方案（80%）", font_size=H1_SIZE, bold=True,
          color=COLOR_H1, space_after=Pt(8))

    if data.get("glossary"):
        add_p(doc, "3.1 名词定义", font_size=H2_SIZE, bold=True, color=COLOR_H2)
        glossary = data["glossary"]
        add_table(doc, ["名词", "定义"],
                  [[r.get("term", ""), r.get("definition", "")]
                   for r in glossary])

    if data.get("er_diagram"):
        add_p(doc, "3.2 E-R图", font_size=H2_SIZE, bold=True, color=COLOR_H2)
        add_p(doc, data["er_diagram"])

    add_p(doc, "3.3 产品结构图（10%）", font_size=H2_SIZE, bold=True, color=COLOR_H2)
    add_p(doc, data.get("structure_diagram", "模块/功能结构…"))

    add_p(doc, "3.4 产品流程图（10%）", font_size=H2_SIZE, bold=True, color=COLOR_H2)
    add_p(doc, data.get("flow_diagram", "多角色泳道图 / 普通流程图…"))

    if data.get("prototype"):
        add_p(doc, "3.5 原型图", font_size=H2_SIZE, bold=True, color=COLOR_H2)
        add_p(doc, data["prototype"])

    # 3.6 需求说明（核心）
    add_p(doc, "3.6 需求说明（60%）", font_size=H2_SIZE, bold=True, color=COLOR_H2)
    requirements = data.get("requirements", [])
    if requirements:
        add_table(doc, ["功能模块", "功能点描述", "优先级"],
                  [[r.get("module", ""), r.get("feature", ""),
                    r.get("priority", "")]
                   for r in requirements])

        for req in requirements:
            if req.get("detail"):
                add_p(doc, "功能点：" + req.get("feature", ""),
                      font_size=H3_SIZE, bold=True, color=COLOR_H2,
                      space_before=Pt(8))
                add_p(doc, req["detail"])

    if data.get("collaboration"):
        add_p(doc, "3.7 对协同方需求", font_size=H2_SIZE, bold=True, color=COLOR_H2)
        add_p(doc, data["collaboration"])

    doc.save(output_path)
    return output_path


def main(argv):
    parser = argparse.ArgumentParser(description="美柚中后台 PRD 生成器")
    parser.add_argument("--output", "-o", required=True, help="输出 .docx 路径")
    parser.add_argument("--data", "-d", help="JSON 数据文件路径")
    args = parser.parse_args(argv)

    if args.data:
        with open(args.data, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.load(io.TextIOWrapper(sys.stdin.buffer, encoding="utf-8"))

    try:
        path = generate_prd(data, args.output)
        print("文档已生成: " + path)
        return 0
    except Exception as e:
        print("生成失败: " + str(e), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
